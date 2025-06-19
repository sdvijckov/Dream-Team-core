# Dream-Team-core/dataset_preparation/src/file_loaders.py
import os
import docx
import re
from .text_cleaner import clean_text # Импортируем из нашего же пакета

def load_paragraphs_from_docx(file_path: str) -> list[str]:
    """
    Загружает текст из файла .docx в виде списка абзацев.
    Каждый абзац предварительно очищается.
    """
    if not os.path.exists(file_path):
        print(f"Ошибка: Файл не найден по пути {file_path}")
        return []
    try:
        doc = docx.Document(file_path)
        # Сохраняем только непустые абзацы после очистки
        paragraphs = [clean_text(para.text) for para in doc.paragraphs if clean_text(para.text)]
        return paragraphs
    except Exception as e:
        print(f"Ошибка при чтении .docx файла {file_path}: {e}")
        return []

def load_paragraphs_from_txt(file_path: str) -> list[str]:
    """
    Загружает текст из файла .txt и разбивает на абзацы.
    Абзацы разделены одной или несколькими пустыми строками.
    Каждый абзац предварительно очищается.
    """
    if not os.path.exists(file_path):
        print(f"Ошибка: Файл не найден по пути {file_path}")
        return []
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        # Разбиваем по двум или более символам новой строки, затем очищаем каждый абзац
        # Также убираем пустые строки, которые могут образоваться после split
        paragraphs = [clean_text(p) for p in re.split(r'\n\s*\n+', content) if clean_text(p)]
        return paragraphs
    except Exception as e:
        print(f"Ошибка при чтении .txt файла {file_path}: {e}")
        return []

if __name__ == '__main__':
    # Для тестирования этого модуля, вам нужно создать тестовые файлы
    # в папке ../input_texts/ относительно текущей папки src
    # Пример: Dream-Team-core/dataset_preparation/input_texts/test_load.docx
    #         Dream-Team-core/dataset_preparation/input_texts/test_load.txt

    current_dir = os.path.dirname(os.path.abspath(__file__))
    input_texts_dir = os.path.join(os.path.dirname(current_dir), "input_texts")
    os.makedirs(input_texts_dir, exist_ok=True) # Создаем папку, если ее нет

    # Тест DOCX
    test_docx_path = os.path.join(input_texts_dir, "test_load.docx")
    if not os.path.exists(test_docx_path):
        print(f"Создание тестового DOCX: {test_docx_path}")
        doc = docx.Document()
        doc.add_paragraph("Первый абзац.  ")
        doc.add_paragraph("  Второй абзац с пробелами.  ")
        doc.add_paragraph("") # Пустой абзац, должен отфильтроваться
        doc.add_paragraph("Третий абзац.")
        doc.save(test_docx_path)

    print(f"\n--- Тестирование DOCX: {test_docx_path} ---")
    docx_paras = load_paragraphs_from_docx(test_docx_path)
    print(f"Загружено абзацев: {len(docx_paras)}")
    for i, p in enumerate(docx_paras):
        print(f"Абзац {i}: '{p}'")

    # Тест TXT
    test_txt_path = os.path.join(input_texts_dir, "test_load.txt")
    if not os.path.exists(test_txt_path):
        print(f"Создание тестового TXT: {test_txt_path}")
        with open(test_txt_path, "w", encoding="utf-8") as f:
            f.write("Текстовый абзац 1.\n\n   \nТекстовый абзац 2 с   пробелами.\n\n\nТекстовый абзац 3.")
    
    print(f"\n--- Тестирование TXT: {test_txt_path} ---")
    txt_paras = load_paragraphs_from_txt(test_txt_path)
    print(f"Загружено абзацев: {len(txt_paras)}")
    for i, p in enumerate(txt_paras):
        print(f"Абзац {i}: '{p}'")